# -*- coding: utf-8 -*-
"""Generate MeterSphere community unlock + org structure construction plan Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from datetime import date


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], 'D9E2F3')
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MeterSphere 社区版解除限制\n与组织架构模块\n完整建设方案')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f'文档版本：V1.0\n编制日期：{date.today().strftime("%Y年%m月%d日")}\n适用范围：MeterSphere V3 自研分支')
    r.font.size = Pt(12)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # 1. Overview
    add_heading(doc, '一、项目概述', 1)
    add_heading(doc, '1.1 建设背景', 2)
    add_para(doc, 'MeterSphere V3 开源社区版在用户数、组织数、资源池等方面存在 License 配额限制；同时企业版扩展点（UserXpackService、LicenseService）未在社区源码中提供完整实现，导致用户管理、多组织等核心能力不可用。')
    add_para(doc, '本项目在 MeterSphere V3 社区版基础上，分两主线推进：')
    add_para(doc, '（1）解除 License 限制：实现社区版 Xpack/License 扩展，恢复多组织、多用户、多资源池能力。', indent=0.5)
    add_para(doc, '（2）组织架构模块：将企业微信通讯录同步为 MeterSphere 组织内只读部门/成员镜像，并提供管理页、同步运维及 Excel 导入扩展。', indent=0.5)

    add_heading(doc, '1.2 总体目标', 2)
    goals = [
        '解除组织数、用户数、资源池等 License 限制，使社区版可支撑企业内部多组织部署；',
        '补齐组织创建、切换、切换选项等后端 API，并完成与前端联调；',
        '建立组织架构数据模型（部门表、同步配置、同步日志、用户扩展字段）；',
        '实现企微通讯录客户端、同步引擎、同步 API 与定时任务；',
        '交付组织架构管理前端、企微配置扩展、Excel 批量导入能力。',
    ]
    for i, g in enumerate(goals, 1):
        add_para(doc, f'{i}. {g}')

    add_heading(doc, '1.3 建设范围', 2)
    add_table(doc, ['范围', '包含', '不包含'], [
        ['License 解除', 'Community Xpack/License 实现、前端 unlimited 开关', '企业版 general-xpack-impl JAR'],
        ['多组织', '组织创建/切换/初始化、前端入口放开', '跨组织数据迁移工具'],
        ['组织架构', '部门树、成员查询、企微同步、Excel 导入', '一人多部门（一期）'],
        ['前端', '组织入口、组织架构管理页、企微配置、导入 UI', '移动端适配'],
    ], [3, 6, 5])

    add_heading(doc, '1.4 工期与人力', 2)
    add_table(doc, ['项', '内容'], [
        ['预估总工期', '约 5–6 周（1 人全职，含联调）'],
        ['阶段划分', 'P0（1 周）→ P1（1 周）→ P2（2 周）→ P3（2 周）'],
        ['已完成', 'task001（2026-06-25）、task002（2026-07-03）'],
        ['进行中/待开始', 'task003 ~ task011'],
    ], [4, 12])

    doc.add_page_break()

    # 2. Architecture
    add_heading(doc, '二、总体架构设计', 1)
    add_heading(doc, '2.1 系统分层', 2)
    add_table(doc, ['层次', '职责', '主要模块'], [
        ['表现层', 'Vue3 + Arco Design，组织入口、管理页、配置页', 'frontend/src/views/setting/'],
        ['接口层', 'REST API，权限校验，操作审计', 'SystemOrganizationController、OrgWecomSyncController 等'],
        ['业务层', '组织管理、用户 Xpack、同步引擎、查询服务', 'OrganizationService、WecomOrgSyncService 等'],
        ['数据层', 'MySQL + Flyway 迁移，MyBatis Mapper', 'organization、department、user、org_sync_log 等'],
        ['外部集成', '企微通讯录 API', 'WecomContactClient'],
    ], [2.5, 4, 7.5])

    add_heading(doc, '2.2 默认产品决策', 2)
    add_table(doc, ['决策项', '推荐默认值'], [
        ['企微配置粒度', '每组织一条（org_wecom_sync_config.organization_id）'],
        ['企微与 MS 组织映射', '一个企微主体 = 一个 MS 组织'],
        ['同步新用户默认角色', 'org_member（组织成员）'],
        ['部门数据来源', '企微同步只读；Excel 导入为 MS 扩展能力'],
        ['管理页可见范围', '系统管理员看全部；组织管理员只看本组织'],
        ['前端 License 开关', 'VITE_MS_UNLIMITED=true（需配合后端 Community Xpack）'],
    ], [5, 11])

    add_heading(doc, '2.3 任务依赖关系', 2)
    add_para(doc, '关键路径：task001 → task002 → task004 → task006 → task007 → task008 → task009')
    add_para(doc, '并行可能：task003 与 task002 后端完成后联调；task010 与 task009 可并行；task011 依赖 task002/004/005。')

    doc.add_page_break()

    # 3. Phase P0
    add_heading(doc, '三、P0 阶段建设方案（第 1 周）', 1)
    add_para(doc, 'P0 目标：解除 License 限制，恢复多组织创建与切换能力，前端入口可用。', bold=True)

    add_heading(doc, '3.1 task001 — 社区版 Xpack 与 License 实现', 2)
    add_para(doc, '状态：已完成（2026-06-25）')
    add_para(doc, '核心工作：')
    items = [
        '新增 CommunityLicenseServiceImpl：validate() 返回 status=valid；',
        '新增 CommunityUserXpackServiceImpl：实现用户增删改落库（user 表 insert、cft_token、user_role_relation）；',
        '新增 CommunityXpackConfiguration：@Configuration + @Bean 注册 License/UserXpack Bean；',
        '回归验证 SimpleUserService、NodeResourcePoolService、TestResourcePoolService 等调用链。',
    ]
    for it in items:
        add_para(doc, f'• {it}', indent=0.5)
    add_para(doc, '关键经验（BUG001）：Community 版 UserXpackService 必须自行完成落库，SimpleUserService 在 xpack 返回 0 后只写日志不 insert；Bean 注册优先使用 @Configuration + @Bean。')

    add_heading(doc, '3.2 task002 — 组织创建与切换 API', 2)
    add_para(doc, '状态：已完成（2026-07-03）')
    add_table(doc, ['API', '方法', '路径', '说明'], [
        ['创建组织', 'POST', '/system/organization/add', '创建组织并完成模板/状态流/字段初始化'],
        ['切换组织', 'POST', '/system/organization/switch', '更新 session 与 last_organization_id'],
        ['切换选项', 'GET', '/system/organization/switch-option', '返回当前用户可切换组织列表'],
    ], [2.5, 1.5, 5.5, 5.5])
    add_para(doc, '核心实现：OrganizationInitService（5 类模板 + Bug 状态流）、OrganizationService.add()、SystemOrganizationLogService.addLog()。')

    add_heading(doc, '3.3 task003 — 前端 License 解除与组织入口', 2)
    add_para(doc, '状态：待开始 | 预估：2 天 | 前置：task001、task002')
    add_table(doc, ['改动点', '文件/配置', '说明'], [
        ['环境变量', 'frontend/.env.development.local.example', 'VITE_MS_UNLIMITED=true'],
        ['License Store', 'store/modules/setting/license.ts', 'hasLicense() 读取 unlimited 开关'],
        ['组织入口', 'organizationAndProject/index.vue、utils.ts', '移除 hasLicense() 拦截'],
        ['资源池', 'resourcePool/index.vue、detail.vue', '多资源池 UI 放开'],
        ['v-xpack 指令', 'directive/validateLicense/index.ts', 'unlimited 时不移除 DOM'],
    ], [2.5, 5.5, 6])

    add_heading(doc, '3.4 P0 里程碑验收（M0）', 2)
    checks = [
        '可创建第 2 个组织并完成初始化（模板、字段、状态流）；',
        '可切换组织上下文，switch-option 返回预期列表；',
        '用户不受 5 人上限限制，资源池不受单池限制；',
        '前端「创建组织」「进入组织」可用（task003 完成后）。',
    ]
    for c in checks:
        add_para(doc, f'☐ {c}')

    doc.add_page_break()

    # 4. Phase P1
    add_heading(doc, '四、P1 阶段建设方案（第 2 周）', 1)
    add_heading(doc, '4.1 task004 — 数据模型与 Flyway 迁移', 2)
    add_para(doc, '状态：待开始 | 预估：1 天 | 建议版本：3.7.0')
    add_table(doc, ['表/变更', '用途'], [
        ['department', '组织部门树，含 wecom_dept_id、parent_id、sync_status'],
        ['org_wecom_sync_config', '每组织企微同步配置（corpId、contactSecret、cron）'],
        ['org_sync_log', '同步执行日志与统计'],
        ['user 扩展字段', 'wecom_userid、department_id、position、sync_status、sync_time'],
    ], [4, 12])
    add_para(doc, '迁移路径：backend/framework/domain/src/main/resources/migration/3.7.0/')

    add_heading(doc, '4.2 task005 — 组织架构查询 API', 2)
    add_para(doc, '状态：待开始 | 预估：2 天')
    add_table(doc, ['API', '路径', '用途'], [
        ['通用部门树', 'GET /department/tree', '业务选人，组织成员可读'],
        ['管理端部门树', 'GET /org-structure/departments/tree', '含 directUserCount、totalUserCount'],
        ['成员分页', 'GET /org-structure/members/page', 'SQL 层分页，禁止内存过滤'],
        ['成员详情', 'GET /org-structure/members/{id}', 'phone/email/wecomUserid 脱敏'],
    ], [3, 5.5, 7.5])
    add_para(doc, '模块：DepartmentQueryService、OrgStructureMemberService、DepartmentController、OrgStructureMemberController。')

    add_heading(doc, '4.3 P1 里程碑验收（M1）', 2)
    for c in ['Flyway 迁移成功；', '部门树/成员分页/成员详情 API 可用；', '可用测试数据验证树形结构与脱敏。']:
        add_para(doc, f'☐ {c}')

    doc.add_page_break()

    # 5. Phase P2
    add_heading(doc, '五、P2 阶段建设方案（第 3–4 周）', 1)
    add_heading(doc, '5.1 task006 — 企微通讯录客户端', 2)
    add_para(doc, '状态：待开始 | 预估：2 天 | 参考：myTapd WecomClient')
    add_table(doc, ['方法', '企微 API', '说明'], [
        ['getAccessToken', 'GET /cgi-bin/gettoken', '内存缓存，过期前 5 分钟刷新'],
        ['listDepartments', 'GET /cgi-bin/department/list', '全量部门'],
        ['listDepartmentUsers', 'GET /cgi-bin/user/list', '部门成员 fetch_child'],
        ['getUser', 'GET /cgi-bin/user/get', '单人详情（登录兜底）'],
    ], [3.5, 4.5, 8])

    add_heading(doc, '5.2 task007 — 组织架构同步引擎', 2)
    add_para(doc, '状态：待开始 | 预估：3 天')
    add_para(doc, '核心原则：先部门后员工；幂等 upsert；空列表跳过失活；空值不覆盖；admin/DEV_ 账号保护。')
    add_para(doc, '部门同步：Pass1 Upsert → Pass2 修正 parent_id → Pass3 失活收敛。')
    add_para(doc, '员工同步：按 wecom_userid upsert；新用户 SimpleUserService.addUser + org_member 绑定。')

    add_heading(doc, '5.3 task008 — 同步 API 与定时任务', 2)
    add_table(doc, ['API', '路径', '说明'], [
        ['手动同步', 'POST /org-wecom/sync/manual', '触发 WecomOrgSyncApplicationService.syncManual()'],
        ['最近状态', 'GET /org-wecom/sync/status', '最近 org_sync_log + last_sync_time'],
        ['日志分页', 'GET /org-wecom/sync/log/page', '历史同步记录'],
        ['保存配置', 'POST /org-wecom/config/save', 'corpId、contactSecret、cron 等'],
    ], [2.5, 5, 6.5])
    add_para(doc, '定时任务：WecomOrgSyncJob（Quartz），组织级 Redis 锁防并发。')

    add_heading(doc, '5.4 P2 里程碑验收（M2）', 2)
    for c in ['手动同步企微通讯录成功；', '部门树与企微一致，用户自动创建并绑定 org_member；', '同步日志、最近状态、定时任务可用。']:
        add_para(doc, f'☐ {c}')

    doc.add_page_break()

    # 6. Phase P3
    add_heading(doc, '六、P3 阶段建设方案（第 5–6 周）', 1)
    add_heading(doc, '6.1 task009 — 组织架构管理前端', 2)
    add_para(doc, '状态：待开始 | 预估：3 天 | 参考：myTapd UserManageView.vue')
    add_para(doc, '布局：左部门树 + 右成员列表 + 详情抽屉 + 同步运维面板。')
    add_para(doc, '路由：/setting/system/org-structure、/setting/organization/org-structure')

    add_heading(doc, '6.2 task010 — 企微配置扩展', 2)
    add_para(doc, '状态：待开始 | 预估：2 天 | 可与 task009 并行')
    add_para(doc, '扩展 corpId、contactSecret（通讯录 Secret）、scheduleCron；GET config、POST test 连通性；Secret 掩码回显。')

    add_heading(doc, '6.3 task011 — Excel 组织架构导入', 2)
    add_para(doc, '状态：待开始 | 预估：3 天 | MS 扩展能力')
    add_table(doc, ['Sheet', '内容'], [
        ['Sheet1 部门', '组织名称、部门路径、上级部门路径、排序'],
        ['Sheet2 用户', '姓名、邮箱、手机、组织名称、部门路径、用户组、职位'],
    ], [3, 13])
    add_para(doc, 'API：GET /org-structure/import/template、POST /org-structure/import')

    add_heading(doc, '6.4 P3 里程碑验收（M3）', 2)
    for c in ['组织架构管理页联调通过；', '企微配置页支持通讯录 Secret + Cron；', 'Excel 导入部门与用户成功。']:
        add_para(doc, f'☐ {c}')

    doc.add_page_break()

    # 7. Data model
    add_heading(doc, '七、数据模型概要', 1)
    add_table(doc, ['实体', '关键字段', '说明'], [
        ['organization', 'id, num, name, enable', 'MS 组织，已有表扩展使用'],
        ['department', 'organization_id, wecom_dept_id, parent_id', '组织内部门树'],
        ['user', 'wecom_userid, department_id, cft_token', '用户表扩展 + 同步字段'],
        ['org_wecom_sync_config', 'corp_id, contact_secret, schedule_cron', '每组织一条企微配置'],
        ['org_sync_log', 'sync_mode, sync_status, dept/user 统计', '同步审计与运维'],
        ['user_role_relation', 'role_id=org_admin/org_member', '组织管理员与成员绑定'],
    ], [3.5, 5, 7.5])

    doc.add_page_break()

    # 8. API summary
    add_heading(doc, '八、API 规划总览', 1)
    add_table(doc, ['分类', '方法', '路径', '任务', '状态'], [
        ['License', 'GET', '/license/validate', 'task001', '已完成'],
        ['用户', 'POST', '/system/user/add', 'task001', '已完成'],
        ['组织', 'POST', '/system/organization/add', 'task002', '已完成'],
        ['组织', 'POST', '/system/organization/switch', 'task002', '已完成'],
        ['组织', 'GET', '/system/organization/switch-option', 'task002', '已完成'],
        ['部门', 'GET', '/department/tree', 'task005', '待开发'],
        ['组织架', 'GET', '/org-structure/departments/tree', 'task005', '待开发'],
        ['组织架', 'GET', '/org-structure/members/page', 'task005', '待开发'],
        ['组织架', 'GET', '/org-structure/members/{id}', 'task005', '待开发'],
        ['同步', 'POST', '/org-wecom/sync/manual', 'task008', '待开发'],
        ['同步', 'GET', '/org-wecom/sync/status', 'task008', '待开发'],
        ['同步', 'GET', '/org-wecom/sync/log/page', 'task008', '待开发'],
        ['配置', 'POST', '/org-wecom/config/save', 'task008/010', '待开发'],
        ['导入', 'POST', '/org-structure/import', 'task011', '待开发'],
    ], [1.8, 1.2, 5.5, 1.5, 1.5])

    doc.add_page_break()

    # 9. Progress
    add_heading(doc, '九、进度跟踪', 1)
    add_table(doc, ['任务', '阶段', '主题', '工期', '状态', '完成日期'], [
        ['task001', 'P0', '社区版 Xpack/License', '2 天', '已完成', '2026-06-25'],
        ['task002', 'P0', '组织创建/切换 API', '3 天', '已完成', '2026-07-03'],
        ['task003', 'P0', '前端 License 解除', '2 天', '待开始', '—'],
        ['task004', 'P1', '数据模型 Flyway', '1 天', '待开始', '—'],
        ['task005', 'P1', '组织架构查询 API', '2 天', '待开始', '—'],
        ['task006', 'P2', '企微通讯录客户端', '2 天', '待开始', '—'],
        ['task007', 'P2', '组织架构同步引擎', '3 天', '待开始', '—'],
        ['task008', 'P2', '同步 API/定时任务', '2 天', '待开始', '—'],
        ['task009', 'P3', '组织架构管理前端', '3 天', '待开始', '—'],
        ['task010', 'P3', '企微配置扩展', '2 天', '待开始', '—'],
        ['task011', 'P3', 'Excel 导入', '3 天', '待开始', '—'],
    ], [1.5, 1, 3.5, 1.2, 1.5, 2.3])

    add_para(doc, '整体进度：11 项任务中 2 项已完成（约 18%），建议按关键路径优先推进 task003 → task004 → task006 → task007 → task008 → task009。')

    doc.add_page_break()

    # 10. Risk
    add_heading(doc, '十、风险与对策', 1)
    add_table(doc, ['风险', '影响', '对策'], [
        ['Xpack Bean 注册失败', '用户创建 NPE、License 无效', '使用 @Configuration + @Bean；参考 BUG001 修复'],
        ['Community 用户未落库', '创建成功但列表无数据', 'UserXpackService 自行 insert user 与 role_relation'],
        ['企微 API 限流/变更', '同步失败', 'Token 缓存、重试、errcode 统一处理'],
        ['空列表误失活', '部门/用户被批量禁用', '空列表保护、内置账号保护'],
        ['与企业版 JAR 冲突', 'Bean 覆盖异常', '@ConditionalOnMissingBean 隔离'],
        ['集成测试依赖 Docker', 'CI/本地测试失败', 'Testcontainers 环境或 Mock 单测补充'],
        ['GPL 协议', '合规风险', '内部部署可接受；保留 MeterSphere Logo/版权'],
    ], [3, 3.5, 9.5])

    doc.add_page_break()

    # 11. Verification
    add_heading(doc, '十一、验证与部署', 1)
    add_heading(doc, '11.1 编译与启动', 2)
    add_para(doc, 'powershell\n.\\mvnw.cmd -f backend\\pom.xml install -pl app -am -DskipTests -DskipAntRunForJenkins=true\n.\\stop.ps1\n.\\start.ps1 -BackendOnly -SkipDeps')

    add_heading(doc, '11.2 关键验证项', 2)
    add_table(doc, ['验证项', '方法', '预期'], [
        ['License', 'GET /license/validate', '{"status":"valid"}'],
        ['用户创建', 'POST /system/user/add', '列表可见新用户'],
        ['组织创建', 'POST /system/organization/add', '返回组织 ID，模板已初始化'],
        ['组织切换', 'POST /system/organization/switch', 'last_organization_id 更新'],
        ['企微同步', 'POST /org-wecom/sync/manual', 'dept/user 统计 SUCCESS'],
    ], [2.5, 4.5, 7])

    add_heading(doc, '11.3 参考资源', 2)
    refs = [
        'docs/task/task000-实施总览与依赖关系.md',
        'docs/summary/community-unlock-and-org-structure.md',
        'docs/summary/组织架构模块设计摘要.md',
        'docs/develop_logs/buglist/2026-06-26-BUG001-创建用户成功但列表不显示.md',
        '参考项目：C:\\SoftWare\\JetBrains\\myTapd（同步引擎、企微客户端）',
        'MeterSphere 论坛 / GitHub Issue #34159',
    ]
    for r in refs:
        add_para(doc, f'• {r}', indent=0.5)

    add_heading(doc, '十二、附录：任务文档索引', 1)
    add_table(doc, ['文档', '路径'], [
        ['task000', 'docs/task/task000-实施总览与依赖关系.md'],
        ['task001', 'docs/task/task001-P0-社区版Xpack与License实现.md'],
        ['task002', 'docs/task/task002-P0-组织创建与切换API.md'],
        ['task003', 'docs/task/task003-P0-前端License解除与组织入口.md'],
        ['task004', 'docs/task/task004-P1-数据模型与Flyway迁移.md'],
        ['task005', 'docs/task/task005-P1-组织架构查询API.md'],
        ['task006', 'docs/task/task006-P2-企微通讯录客户端.md'],
        ['task007', 'docs/task/task007-P2-组织架构同步引擎.md'],
        ['task008', 'docs/task/task008-P2-同步API与定时任务.md'],
        ['task009', 'docs/task/task009-P3-组织架构管理前端.md'],
        ['task010', 'docs/task/task010-P3-企微配置扩展.md'],
        ['task011', 'docs/task/task011-P3-Excel组织架构导入.md'],
    ], [2, 14])

    return doc


if __name__ == '__main__':
    output = r'C:\SoftWare\JetBrains\metersphere\docs\task\MeterSphere-Community-Unlock-OrgStructure-Construction-Plan.docx'
    build_document().save(output)
    print(f'Generated: {output}')
