# -*- coding: utf-8 -*-
"""Generate Word document for Test Data Center + WeCom Bot plan."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_shading(cell, color: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9E2F3")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("MeterSphere 测试数据中心\n+\n企业微信助手\n综合建设方案")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.name = "Microsoft YaHei"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("版本：V1.0    日期：2026年7月")
    sr.font.size = Pt(11)
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph()

    # 1
    add_heading(doc, "1. 背景与目标", 1)
    add_heading(doc, "1.1 业务背景", 2)
    add_para(doc, "测试团队需要在 MeterSphere 平台中集中管理业务测试数据（如二维码链接、流水号等），供自动化执行与平台使用者共同使用。同时，希望通过企业微信 @「Metersphere助手」实现交互式取数：用户说明需求 → 平台提取并打包 → 回发到企业微信。")

    add_heading(doc, "1.2 建设目标", 2)
    add_table(doc, ["目标", "说明"], [
        ["统一存储", "一处维护，多场景复用"],
        ["自动化可用", "执行引擎通过变量语法引用"],
        ["人工可用", "Web 界面 + 企业微信双通道领取"],
        ["可控消费", "流水号等支持领取锁定，避免重复使用"],
        ["企微交互", "@助手 → 推送 → 取数 → 打包 → 回复"],
    ], [1.2, 4.5])

    # 2
    add_heading(doc, "2. 现状分析", 1)
    add_heading(doc, "2.1 MeterSphere 现有能力", 2)
    add_table(doc, ["能力", "位置", "自动化", "人工", "结论"], [
        ["全局参数", "项目管理→环境管理→全局参数", "✓", "△", "仅适合少量固定值"],
        ["环境变量", "环境管理→各环境", "✓", "△", "按环境隔离，非数据中心"],
        ["场景 CSV 变量", "接口测试→场景→变量", "✓", "✗", "绑定场景，分散管理"],
        ["环境「数据源」", "环境管理→数据库", "✓ SQL", "✗", "JDBC 连接，非测试数据"],
        ["文件管理", "项目管理→文件管理", "✓", "△", "无语义、无状态"],
        ["自定义函数", "项目管理", "✓", "✗", "适合生成，不适合存储"],
        ["自定义字段", "用例/缺陷模板", "✗", "✓", "用例元数据，非执行数据"],
    ])

    add_para(doc, "结论：目前没有专门的「测试数据中心」模块。现有变量模型（CommonVariables）支持 CONSTANT / LIST / JSON 三种类型，已在 API 执行链路中注入 JMeter，但无法满足数据集管理、消费锁定、跨模块复用、企微交互取数等需求。", bold=True)

    add_heading(doc, "2.2 企业微信现有能力", 2)
    add_table(doc, ["能力", "用途", "能否 @ 交互"], [
        ["消息管理→Webhook 机器人", "测试报告等事件通知（出站）", "否"],
        ["系统设置→扫码登录", "SSO 登录（corpId/agentId/appSecret）", "否"],
    ], [2.5, 2.5, 1.5])

    add_para(doc, "现有 WeComClient 仅支持向 Webhook POST 文本，无入站消息回调，不能实现 @ 助手交互。企微助手需全新建设，与现有通知机器人职责分离。", bold=True)

    # 3
    add_heading(doc, "3. 总体架构", 1)
    add_para(doc, "核心原则：")
    add_bullet(doc, "测试数据中心是唯一数据来源（Single Source of Truth）")
    add_bullet(doc, "Web、自动化、企微三条通道共用同一套数据服务与权限模型")
    add_bullet(doc, "企微通知 Webhook 与 Metersphere助手 分开配置、分开开发")

    add_para(doc, "架构示意：")
    add_code_block(doc, """使用者
  ├── 测试人员（Web）────────────→ 测试数据中心 Web
  ├── 自动化执行 ────────────────→ 执行引擎 TestDataResolver
  └── 测试人员（企业微信）
        │ @Metersphere助手
        ▼
      企业微信自建应用
        │ POST 推送消息
        ▼
      MeterSphere 企微回调 /wecom/bot/callback
        │ 解析 → 取数 → 打包
        ▼
      测试数据服务 ←→ test_data_pool / set / record / log
        │ message/send
        ▼
      企业微信 → 用户收到打包数据""")

    # 4
    add_heading(doc, "4. 模块一：测试数据中心", 1)
    add_heading(doc, "4.1 领域模型", 2)
    add_table(doc, ["实体", "说明", "示例"], [
        ["数据池 Pool", "项目级容器", "「营销活动测试数据」"],
        ["数据集 Set", "按类型分组，有唯一 code", "qr_link（二维码）、serial_no（流水号）"],
        ["数据记录 Record", "单条数据 + 状态", "AVAILABLE / LOCKED / CONSUMED / EXPIRED"],
        ["消费日志 Log", "审计追踪", "记录操作人、来源（Web/自动化/企微）"],
    ])

    add_heading(doc, "4.2 后端设计", 2)
    add_table(doc, ["项", "建议"], [
        ["所属模块", "project-management 或独立 test-data-management"],
        ["权限", "PROJECT_TEST_DATA:READ / CREATE / UPDATE / DELETE / CONSUME"],
        ["核心 API", "CRUD、批量导入、领取（consume）、释放（release）、按条件查询"],
    ], [1.5, 4.2])

    add_para(doc, "与执行引擎集成 — 变量语法：")
    add_table(doc, ["语法", "行为"], [
        ["${testData.qr_link.next}", "领取下一条可用数据并锁定"],
        ["${testData.serial_no.peek}", "只读预览，不消费"],
        ["${testData.account.fixed}", "读取固定可用数据"],
    ], [2.5, 3.2])

    add_heading(doc, "4.3 前端设计", 2)
    add_para(doc, "路径：项目管理 → 测试数据中心")
    add_bullet(doc, "数据池列表：按项目展示，搜索、新建")
    add_bullet(doc, "数据集详情：表格展示记录，支持 Excel/CSV 导入")
    add_bullet(doc, "单条操作：复制、标记已用、释放、设过期")
    add_bullet(doc, "领取面板：测试人员一键领取可用数据")
    add_bullet(doc, "引用说明：展示 ${testData.xxx} 及企微指令用法")

    add_heading(doc, "4.4 过渡方案", 2)
    add_table(doc, ["阶段", "方案", "适用"], [
        ["临时", "全局参数/环境变量（CONSTANT、LIST）", "少量固定值，无消费锁定"],
        ["过渡", "文件管理 + CSV + 场景变量", "批量 API 驱动，无人工管理"],
        ["正式", "测试数据中心（本方案）", "完整需求"],
    ])

    # 5
    add_heading(doc, "5. 模块二：企业微信 Metersphere助手", 1)
    add_heading(doc, "5.1 交互流程", 2)
    add_table(doc, ["步骤", "用户描述", "技术实现"], [
        ["1", "用户 @Metersphere助手 输出消息", "企业微信捕获并推送"],
        ["2", "通过接口向 MeterSphere 推送消息", "POST /wecom/bot/callback"],
        ["3", "MeterSphere 根据消息内容提取数据", "意图解析 → TestDataFetchService"],
        ["4", "将数据打包", "TestDataPackService"],
        ["5", "通过接口向企业微信发送打包数据", "WeComAppClient.message/send"],
    ], [0.5, 2.2, 3.0])

    add_para(doc, "完整时序：")
    add_code_block(doc, """测试人员 → 企业微信：@Metersphere助手 项目=营销活动 二维码*2 流水号*5
企业微信 → MeterSphere：POST /wecom/bot/callback（加密消息）
MeterSphere：验签解密 → 幂等检查 → 解析指令 → 用户鉴权 → 立即返回 success
MeterSphere → 测试数据服务：按类型领取数据（异步）
MeterSphere：打包（文本 / Excel）
MeterSphere → 企业微信：POST message/send（应用 API）
企业微信 → 测试人员：群聊/单聊展示结果""")

    add_heading(doc, "5.2 企业微信侧配置", 2)
    add_bullet(doc, "创建自建应用「Metersphere助手」，记录 corpId、agentId、Secret")
    add_bullet(doc, "开启「接收消息」，配置回调 URL、Token、EncodingAESKey")
    add_bullet(doc, "回调 URL：https://{ms-domain}/wecom/bot/callback（HTTPS 公网）")
    add_bullet(doc, "将应用加入目标群或开放单聊")
    add_para(doc, "注意：现有「消息管理→Webhook 机器人」仅支持出站通知，不能接收 @ 消息，不可复用。", bold=True)

    add_heading(doc, "5.3 MeterSphere 入站接口", 2)
    add_code_block(doc, """GET  /wecom/bot/callback   # URL 验证（配置回调时）
POST /wecom/bot/callback   # 接收用户消息推送""")
    add_table(doc, ["处理要点", "说明"], [
        ["验签解密", "标准企业微信加解密协议"],
        ["快速 ACK", "5 秒内返回 success，取数异步执行"],
        ["幂等", "MsgId + Redis 去重"],
        ["过滤", "仅处理 @ 本应用且为文本的消息"],
    ])

    add_heading(doc, "5.4 消息解析（第一期：结构化指令）", 2)
    add_para(doc, "用户示例：@Metersphere助手 项目=营销活动 二维码链接*2 流水号*5")
    add_code_block(doc, """解析结果：
{
  "projectCode": "营销活动",
  "items": [
    { "type": "qr_link", "count": 2 },
    { "type": "serial_no", "count": 5 }
  ],
  "mode": "CONSUME"
}""")
    add_table(doc, ["用户说法", "数据集 code"], [
        ["二维码 / 二维码链接", "qr_link"],
        ["流水号", "serial_no"],
    ], [2.5, 2.5])

    add_heading(doc, "5.5 数据打包与发送", 2)
    add_table(doc, ["数据量", "打包方式", "企微消息类型"], [
        ["少量（≤10 条）", "Markdown / 文本摘要", "msgtype: markdown 或 text"],
        ["较多（10–100 条）", "生成 Excel", "上传临时素材 → msgtype: file"],
        ["含敏感信息", "群聊摘要 + 单聊详情", "两次 message/send"],
    ])

    add_para(doc, "回复示例：")
    add_code_block(doc, """【MeterSphere 测试数据】
项目：营销活动 | 领取人：张三 | 时间：2025-06-25 15:00

▸ 二维码链接（2条）
  1. https://qr.example.com/a1
  2. https://qr.example.com/a2

▸ 流水号（5条）
  SN001, SN002, SN003, SN004, SN005""")

    add_heading(doc, "5.6 MeterSphere 出站 API", 2)
    add_para(doc, "新建 WeComAppClient（与现有 Webhook 版 WeComClient 并存）：")
    add_table(doc, ["步骤", "企业微信 API"], [
        ["获取 token", "GET /cgi-bin/gettoken?corpid=&corpsecret="],
        ["发送消息", "POST /cgi-bin/message/send?access_token="],
        ["上传文件", "POST /cgi-bin/media/upload"],
    ], [1.5, 4.2])

    add_heading(doc, "5.7 身份与权限", 2)
    add_code_block(doc, """企微 FromUserName → wecom_user_binding 表 → ms_user_id
  → 项目成员 + PROJECT_TEST_DATA:CONSUME""")
    add_para(doc, "绑定方式：① 扫码登录时自动绑定  ② 管理后台手动绑定  ③ 首次 @ 时验证码确认")

    add_heading(doc, "5.8 平台配置项", 2)
    add_table(doc, ["配置", "用途"], [
        ["启用开关", "是否开启 Metersphere助手"],
        ["corpId / agentId / secret", "应用凭证"],
        ["callbackToken / encodingAesKey", "入站加解密"],
        ["defaultProjectId", "默认项目"],
        ["dataTypeAliases", "关键词 → 数据集映射"],
        ["maxFetchPerRequest", "单次领取上限（如 20 条）"],
        ["sensitiveDataPolicy", "群聊脱敏 / 强制单聊"],
        ["allowedChatIds", "允许使用的群白名单（可选）"],
    ])

    # 6
    add_heading(doc, "6. 代码结构规划", 1)
    add_code_block(doc, """backend/
├── test-data/                          # 测试数据中心
│   ├── domain/                         # pool / set / record / log
│   ├── service/
│   │   ├── TestDataPoolService.java
│   │   ├── TestDataConsumeService.java
│   │   └── TestDataPackService.java
│   ├── controller/TestDataController.java
│   └── resolver/TestDataResolver.java
│
├── wecom-bot/                          # 企微助手
│   ├── controller/WeComBotCallbackController.java
│   ├── service/
│   │   ├── WeComCryptoService.java
│   │   ├── WeComBotOrchestrator.java
│   │   ├── WeComMessageParser.java
│   │   ├── WeComAppClient.java
│   │   └── WeComUserBindingService.java
│
frontend/
├── views/project-management/testDataCenter/
└── views/setting/system/config/weComBot/""")

    # 7
    add_heading(doc, "7. 与现有模块关系", 1)
    add_table(doc, ["现有模块", "关系", "处理方式"], [
        ["全局参数/环境变量", "临时数据源", "过渡期可用，正式后迁移"],
        ["消息管理 Webhook 机器人", "事件通知（出站）", "保留，职责不变"],
        ["企业微信扫码登录", "SSO + 用户身份", "复用 corpId 凭证与 userid 绑定"],
        ["WeComClient（Webhook）", "通知发送", "保留，新建 WeComAppClient"],
        ["文件管理", "CSV 文件存储", "数据中心导入可复用"],
    ])

    # 8
    add_heading(doc, "8. 安全与运维", 1)
    add_table(doc, ["项", "要求"], [
        ["公网 HTTPS", "企微回调必须可达；内网部署需网关或 Bot 网关"],
        ["消息幂等", "MsgId 去重，防止重复领取"],
        ["异步处理", "回调先 ACK，取数打包异步完成"],
        ["权限校验", "每次领取校验平台用户 + 项目权限"],
        ["敏感数据", "流水号等支持群聊脱敏、详情单聊"],
        ["限流", "按人/群限制领取频率"],
        ["审计", "记录 msgId、userid、原始指令、领取记录、发送结果"],
    ])

    # 9
    add_heading(doc, "9. 实施计划", 1)
    add_heading(doc, "9.1 分期路线", 2)
    add_table(doc, ["阶段", "内容", "依赖", "预估"], [
        ["P0", "测试数据中心：表结构+CRUD+Web 管理页+人工复制", "无", "2–3 周"],
        ["P1", "执行引擎集成 ${testData.*} + 消费锁定", "P0", "2 周"],
        ["P2", "企微应用配置+回调验签+能收能发", "公网 HTTPS", "1 周"],
        ["P3", "消息解析+用户绑定+对接数据中心+文本回复", "P0+P2", "2 周"],
        ["P4", "Excel 打包+文件消息+消费日志+限流脱敏", "P3", "1–2 周"],
        ["P5", "功能用例/UI 自动化统一引用；NL/卡片交互（可选）", "P1", "2–3 周"],
    ])

    add_heading(doc, "9.2 里程碑", 2)
    add_bullet(doc, "P0+P1 → 平台内可用的测试数据中心（Web + 自动化）")
    add_bullet(doc, "P2+P3 → 企微 @助手 可取数并文本回复")
    add_bullet(doc, "P4 → 生产级企微交互（批量打包、审计、安全）")
    add_bullet(doc, "P5 → 全场景覆盖 + 体验优化")

    # 10
    add_heading(doc, "10. 方案对比与选型结论", 1)
    add_heading(doc, "10.1 测试数据存储", 2)
    add_table(doc, ["方案", "成本", "适用", "结论"], [
        ["复用全局参数", "无", "少量固定值", "仅作过渡"],
        ["文件管理+CSV", "低", "API 批量驱动", "可作过渡"],
        ["新建测试数据中心", "中高", "完整需求", "推荐"],
    ])

    add_heading(doc, "10.2 企微接入", 2)
    add_table(doc, ["方案", "适用", "结论"], [
        ["复用 Webhook 机器人", "仅通知", "不可行（无入站）"],
        ["自建应用+回调入站+API 出站", "@ 交互取数", "推荐"],
        ["独立 Bot 网关", "内网部署/多系统", "无公网时备选"],
    ])

    add_heading(doc, "10.3 最终选型", 2)
    add_table(doc, ["模块", "选型"], [
        ["数据存储", "新建「测试数据中心」模块"],
        ["自动化引用", "${testData.{code}.next/peek/fixed}"],
        ["人工 Web 领取", "测试数据中心 Web 界面"],
        ["企微交互", "新建「Metersphere助手」自建应用"],
        ["现有通知机器人", "保留，与助手分离"],
    ])

    # 11
    add_heading(doc, "11. 待确认事项", 1)
    add_bullet(doc, "部署环境：MeterSphere 是否有公网 HTTPS？")
    add_bullet(doc, "交互场景：助手主要用于群聊、单聊，还是两者兼有？")
    add_bullet(doc, "消费策略：@ 领取后是否立即锁定为「已用」？")
    add_bullet(doc, "数据量级：单次领取通常多少条？")
    add_bullet(doc, "敏感数据：是否要求群聊脱敏、详情仅单聊可见？")
    add_bullet(doc, "企微应用：助手与扫码登录共用 agentId 还是独立应用？")
    add_bullet(doc, "数据类型清单：首批需支持哪些类型及别名？")

    # 12
    add_heading(doc, "12. 方案价值总结", 1)
    add_para(doc, "本方案在 MeterSphere 中新增两大能力，形成完整闭环：")
    add_code_block(doc, """                    ┌─────────────────────┐
                    │   测试数据中心       │
                    │  （统一数据底座）     │
                    └─────────┬───────────┘
                              │
          ┌───────────────────┼───────────────────┐
          │                   │                   │
    ┌─────▼─────┐     ┌──────▼──────┐    ┌──────▼──────┐
    │ Web 界面   │     │ 自动化执行   │    │ 企微助手     │
    │ 浏览/领取  │     │ ${testData.*}│    │ @助手取数    │
    └───────────┘     └─────────────┘    └─────────────┘""")
    add_bullet(doc, "测试人员：Web 或企业微信均可领取，体验一致")
    add_bullet(doc, "自动化：执行时自动引用，支持消费锁定")
    add_bullet(doc, "管理员：集中维护、导入、审计，权限可控")
    add_bullet(doc, "平台：不破坏现有通知与 SSO 能力，增量建设、职责清晰")

    doc.add_paragraph()
    add_para(doc, "— 文档结束 —")

    return doc


def main():
    out_dir = Path(__file__).resolve().parent
    filename = "MeterSphere-TestDataCenter-WeComBot-Plan.docx"
    cn_filename = "\u6d4b\u8bd5\u6570\u636e\u4e2d\u5fc3\u4e0e\u4f01\u4e1a\u5fae\u4fe1\u52a9\u624b\u7efc\u5408\u5efa\u8bbe\u65b9\u6848.docx"
    out_path = out_dir / filename
    cn_path = out_dir / cn_filename
    doc = build_document()
    doc.save(str(out_path))
    doc.save(str(cn_path))
    print(f"Generated: {out_path}")
    print(f"Generated: {cn_path}")


if __name__ == "__main__":
    main()
